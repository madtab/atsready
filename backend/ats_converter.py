#!/usr/bin/env python3
"""
ATS Resume Converter — "Black Box" Engine
==========================================
Takes a PDF resume, extracts content, scores ATS-readability,
and outputs a clean, ATS-optimised .docx file.

Dependencies:
    pip install pdfplumber python-docx

Usage:
    python ats_converter.py input_resume.pdf output_resume.docx

Returns exit code 0 on success, 1 on failure.
Also prints a JSON summary to stdout with the ATS score and issues found.
"""

import sys
import json
import re
import os
from pathlib import Path

import pdfplumber
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ─── ATS Scoring Rules ────────────────────────────────────────────────────────

ATS_PENALTIES = {
    "no_email": ("No email address found", -15),
    "no_phone": ("No phone number found", -10),
    "uses_tables": ("PDF appears to use tables/columns (ATS parsers struggle with these)", -20),
    "too_short": ("Very little text extracted — may be image-based PDF", -30),
    "special_chars": ("Excessive special characters or symbols detected", -10),
    "no_sections": ("Could not detect standard section headings", -15),
    "too_long": ("Resume exceeds 2 pages of content", -5),
}

SECTION_HEADINGS = [
    "experience", "work experience", "employment", "professional experience",
    "education", "qualifications", "academic",
    "skills", "technical skills", "core competencies", "key skills",
    "summary", "profile", "personal statement", "objective", "about me", "about",
    "certifications", "certificates", "training",
    "projects", "portfolio",
    "achievements", "awards", "honours", "honors",
    "references", "referees",
    "languages", "interests", "hobbies", "volunteer", "volunteering",
    "publications", "research",
    "contact", "contact information", "personal details",
]


def normalize_text(text: str) -> str:
    """
    Normalize Unicode characters that cause matching issues.
    E.g. Turkish İ (U+0130) and ı (U+0131) to standard I/i.
    """
    import unicodedata
    # NFKD decomposition splits combined chars, then strip combining marks
    normalized = unicodedata.normalize('NFKD', text)
    # Remove combining marks (accents, dots above, etc.)
    result = ''.join(
        c for c in normalized
        if not unicodedata.combining(c)
    )
    # Collapse multiple spaces into single space
    result = re.sub(r'  +', ' ', result)
    return result


def collapse_spaced_text(text: str) -> str:
    """
    Collapse decoratively spaced text like 'E D U C A T I O N' → 'EDUCATION'.
    Also handles mixed case like 'W o r k  E x p e r i e n c e'.
    """
    import re

    def _collapse_match(match):
        spaced = match.group(0)
        collapsed = spaced.replace(" ", "")
        # Only collapse if result looks like a real word (3+ chars)
        if len(collapsed) >= 3:
            return collapsed
        return spaced

    # Pattern: single chars separated by spaces (at least 3 chars worth)
    # e.g. "E D U C A T I O N" or "S K İ L L S"
    text = re.sub(
        r'\b(?:[A-ZİŞÇÖÜĞa-zışçöüğ] ){2,}[A-ZİŞÇÖÜĞa-zışçöüğ]\b',
        _collapse_match,
        text
    )
    return text


def clean_section_prefix(text: str) -> str:
    """Remove decorative section prefixes like '//' or '▪' from headings."""
    import re
    # Remove leading // or similar decorative chars
    text = re.sub(r'^[/\|•▪►■●]+\s*', '', text.strip())
    return text.strip()


def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extract text from a PDF with column-aware extraction.
    Detects two-column layouts and extracts left then right column
    to avoid interleaving content from different columns.
    """
    full_text = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            words = page.extract_words(keep_blank_chars=True)
            if not words:
                continue

            page_width = page.width

            # Use body words (skip top header area) for column detection
            # Headers often span full width and confuse the analysis
            body_y_start = page.height * 0.2
            body_words = [w for w in words if w["top"] > body_y_start]

            if len(body_words) < 10:
                body_words = words  # fallback if very little body content

            # Find the column gap by analysing x-positions of body words
            x_starts = sorted(set(round(w["x0"]) for w in body_words))

            # Find the largest horizontal gap in the middle portion of the page
            best_gap_pos = None
            max_gap_size = 0

            for i in range(1, len(x_starts)):
                gap_size = x_starts[i] - x_starts[i-1]
                gap_center = (x_starts[i] + x_starts[i-1]) / 2
                # Look for gaps in the middle 60% of the page
                if (gap_size > max_gap_size and
                        page_width * 0.20 < gap_center < page_width * 0.80):
                    max_gap_size = gap_size
                    best_gap_pos = gap_center

            # It's two-column if there's a clear gap (at least 20px)
            # and both sides have substantial content
            is_two_column = False
            col_boundary = page_width / 2

            if best_gap_pos and max_gap_size > 20:
                col_boundary = best_gap_pos
                left_count = len([w for w in body_words if w["x0"] < col_boundary])
                right_count = len([w for w in body_words if w["x0"] >= col_boundary])

                if left_count > 5 and right_count > 5:
                    is_two_column = True

            if is_two_column:
                # Extract header area (top 20%) separately without column split
                header_words = sorted(
                    [w for w in words if w["top"] <= body_y_start],
                    key=lambda w: (round(w["top"], 1), w["x0"])
                )
                header_text = _words_to_text(header_words) if header_words else ""

                # Split body words using detected boundary
                body_left = sorted(
                    [w for w in words if w["top"] > body_y_start and w["x0"] < col_boundary],
                    key=lambda w: (round(w["top"], 1), w["x0"])
                )
                body_right = sorted(
                    [w for w in words if w["top"] > body_y_start and w["x0"] >= col_boundary],
                    key=lambda w: (round(w["top"], 1), w["x0"])
                )

                # Build text from each column
                left_text = _words_to_text(body_left)
                right_text = _words_to_text(body_right)

                # Main content column usually has more text — put it first
                if len(right_text) >= len(left_text):
                    body_text = right_text + "\n\n" + left_text
                else:
                    body_text = left_text + "\n\n" + right_text

                page_text = header_text + "\n\n" + body_text if header_text else body_text
            else:
                # Single column — use default extraction
                text = page.extract_text()
                page_text = text if text else ""

            if page_text.strip():
                full_text.append(page_text)

    result = "\n".join(full_text)

    # Post-process: collapse spaced-out decorative text and normalize Unicode
    result = collapse_spaced_text(result)
    result = normalize_text(result)

    return result


def _words_to_text(words: list) -> str:
    """Convert a list of word dicts (with position info) into readable text."""
    if not words:
        return ""

    lines = []
    current_line_words = []
    current_top = None
    line_threshold = 3  # pixels — words within this vertical gap are same line

    for w in words:
        if current_top is None:
            current_top = w["top"]
            current_line_words.append(w)
        elif abs(w["top"] - current_top) <= line_threshold:
            current_line_words.append(w)
        else:
            # New line — flush current
            line_text = " ".join(
                word["text"] for word in
                sorted(current_line_words, key=lambda x: x["x0"])
            )
            lines.append(line_text)
            current_line_words = [w]
            current_top = w["top"]

    # Flush last line
    if current_line_words:
        line_text = " ".join(
            word["text"] for word in
            sorted(current_line_words, key=lambda x: x["x0"])
        )
        lines.append(line_text)

    return "\n".join(lines)


def _is_contact_line(line: str) -> bool:
    """Check if a line looks like contact information."""
    import re
    stripped = line.strip()
    # Phone number
    if re.search(r'[\+]?[\d\s\-\(\)]{7,15}$', stripped):
        return True
    # Email
    if re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', stripped):
        return True
    # URL / website
    if re.search(r'www\.|https?://|\.com|\.co\.uk|\.org|\.net', stripped, re.IGNORECASE):
        return True
    # Address-like (number + street name pattern)
    if re.search(r'^\d+\s+\w+\s+(St|Street|Rd|Road|Ave|Avenue|Dr|Drive|Ln|Lane|Blvd|Way|Anywhere)', stripped, re.IGNORECASE):
        return True
    return False


def detect_sections(text: str) -> list[dict]:
    """
    Detect resume sections by matching known heading patterns.
    Handles decorative prefixes (like //) and spaced-out headings.
    Also detects orphaned contact info and creates a Contact section.
    Returns a list of {"heading": str, "content": str} dicts.
    """
    lines = text.split("\n")
    sections = []
    current_heading = "HEADER"
    current_lines = []

    for line in lines:
        stripped = line.strip()

        # Clean decorative prefixes before matching
        cleaned = clean_section_prefix(stripped)
        lower = cleaned.lower().rstrip(":")

        # Check if this line is a section heading
        is_heading = False
        for heading in SECTION_HEADINGS:
            if lower == heading or lower.startswith(heading + " "):
                is_heading = True
                matched_heading = cleaned
                break

        if is_heading:
            # Save previous section
            if current_lines:
                sections.append({
                    "heading": current_heading,
                    "content": "\n".join(current_lines).strip()
                })
            current_heading = matched_heading
            current_lines = []
        else:
            if stripped:
                current_lines.append(stripped)

    # Don't forget last section
    if current_lines:
        sections.append({
            "heading": current_heading,
            "content": "\n".join(current_lines).strip()
        })

    # Post-process: find orphaned contact info blocks inside non-contact sections
    # and split them into a separate Contact section
    final_sections = []
    for sec in sections:
        heading = sec["heading"]
        content_lines = sec["content"].split("\n")

        # Skip if this is already a contact section or HEADER
        if heading.lower() in ("contact", "contact information", "personal details", "HEADER".lower()):
            final_sections.append(sec)
            continue

        # Look for a run of contact-like lines within this section
        contact_start = None
        contact_end = None
        for i, line in enumerate(content_lines):
            if _is_contact_line(line):
                if contact_start is None:
                    contact_start = i
                contact_end = i
            else:
                # If we had a contact run going and hit a non-contact line,
                # only break if we found at least 2 contact lines
                if contact_start is not None and (contact_end - contact_start) >= 1:
                    break
                contact_start = None
                contact_end = None

        # If we found a block of 2+ contact lines, split them out
        if contact_start is not None and contact_end is not None and (contact_end - contact_start) >= 1:
            before = content_lines[:contact_start]
            contact = content_lines[contact_start:contact_end + 1]
            after = content_lines[contact_end + 1:]

            # Add section content before the contact block
            if before:
                final_sections.append({
                    "heading": heading,
                    "content": "\n".join(before).strip()
                })

            # Add the contact section
            final_sections.append({
                "heading": "Contact",
                "content": "\n".join(contact).strip()
            })

            # Add remaining content after the contact block
            if after:
                remaining = "\n".join(after).strip()
                if remaining:
                    final_sections.append({
                        "heading": heading + " (continued)",
                        "content": remaining
                    })
        else:
            final_sections.append(sec)

    return final_sections


def score_ats_readability(text: str, sections: list[dict]) -> dict:
    """
    Score the resume for ATS compatibility (0-100).
    Returns {"score": int, "issues": [str], "passed": [str]}.
    """
    score = 100
    issues = []
    passed = []

    # Check for email
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    if re.search(email_pattern, text):
        passed.append("Email address detected")
    else:
        score += ATS_PENALTIES["no_email"][1]
        issues.append(ATS_PENALTIES["no_email"][0])

    # Check for phone
    phone_pattern = r'[\+]?[\d\s\-\(\)]{7,15}'
    if re.search(phone_pattern, text):
        passed.append("Phone number detected")
    else:
        score += ATS_PENALTIES["no_phone"][1]
        issues.append(ATS_PENALTIES["no_phone"][0])

    # Check text length (image-based PDFs extract very little)
    word_count = len(text.split())
    if word_count < 50:
        score += ATS_PENALTIES["too_short"][1]
        issues.append(ATS_PENALTIES["too_short"][0])
    else:
        passed.append(f"Text extraction successful ({word_count} words)")

    # Check for too many special characters (indicates decorative PDFs)
    special_count = len(re.findall(r'[■●►▪◆★☆•→←↑↓▶◀⬛⬜🔹🔸]', text))
    if special_count > 10:
        score += ATS_PENALTIES["special_chars"][1]
        issues.append(ATS_PENALTIES["special_chars"][0])
    else:
        passed.append("Clean character usage")

    # Check for standard section headings
    non_header_sections = [s for s in sections if s["heading"] != "HEADER"]
    if len(non_header_sections) >= 2:
        headings_found = [s["heading"] for s in non_header_sections]
        passed.append(f"Section headings detected: {', '.join(headings_found[:5])}")
    else:
        score += ATS_PENALTIES["no_sections"][1]
        issues.append(ATS_PENALTIES["no_sections"][0])

    # Check length (proxy via word count — ~500 words/page)
    if word_count > 1200:
        score += ATS_PENALTIES["too_long"][1]
        issues.append(ATS_PENALTIES["too_long"][0])
    else:
        passed.append("Appropriate length")

    return {
        "score": max(0, min(100, score)),
        "issues": issues,
        "passed": passed,
        "word_count": word_count,
    }


def build_ats_docx(sections: list[dict], output_path: str, score_data: dict):
    """
    Build a clean, ATS-optimised .docx file.
    Uses only simple formatting that ATS systems can reliably parse:
    - Single column layout
    - Standard fonts (Calibri)
    - Clear section headings
    - No tables, text boxes, or graphics
    - Consistent heading hierarchy
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Set narrow margins for more content space
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    for i, sec in enumerate(sections):
        heading_text = sec["heading"]
        content_text = sec["content"]

        if heading_text == "HEADER":
            # First section is the name/contact block
            lines = content_text.split("\n")
            if lines:
                # Name line — large and bold
                name_para = doc.add_paragraph()
                name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                name_run = name_para.add_run(lines[0])
                name_run.bold = True
                name_run.font.size = Pt(18)
                name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

                # Contact details
                for line in lines[1:]:
                    contact_para = doc.add_paragraph()
                    contact_para.paragraph_format.space_before = Pt(1)
                    contact_para.paragraph_format.space_after = Pt(1)
                    contact_run = contact_para.add_run(line)
                    contact_run.font.size = Pt(10)
                    contact_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

                # Add a thin line separator
                separator = doc.add_paragraph()
                separator.paragraph_format.space_before = Pt(6)
                separator.paragraph_format.space_after = Pt(6)
                sep_run = separator.add_run("─" * 70)
                sep_run.font.size = Pt(6)
                sep_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        else:
            # Section heading
            heading_para = doc.add_paragraph()
            heading_para.paragraph_format.space_before = Pt(14)
            heading_para.paragraph_format.space_after = Pt(4)
            heading_run = heading_para.add_run(heading_text.upper())
            heading_run.bold = True
            heading_run.font.size = Pt(12)
            heading_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

            # Subtle underline effect
            underline_para = doc.add_paragraph()
            underline_para.paragraph_format.space_before = Pt(0)
            underline_para.paragraph_format.space_after = Pt(6)
            ul_run = underline_para.add_run("─" * 70)
            ul_run.font.size = Pt(4)
            ul_run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)

            # Section content — preserve line breaks as paragraphs
            # Add extra spacing before new role entries to visually separate them
            content_lines = content_text.split("\n")
            non_empty = [(i, l) for i, l in enumerate(content_lines) if l.strip()]

            # Pre-scan: find indices of lines that contain date ranges
            date_pattern = re.compile(
                r'\b(19|20)\d{2}\s*[-–—]\s*((19|20)\d{2}|present|now|current)\b',
                re.IGNORECASE
            )
            date_indices = set()
            for i, line in enumerate(content_lines):
                if date_pattern.search(line):
                    date_indices.add(i)

            # Find "entry start" lines: the job title is typically 2 lines
            # before the date (title, company, date pattern)
            entry_start_indices = set()
            for di in date_indices:
                # Walk back up to 2 non-empty lines before the date
                candidates = []
                j = di - 1
                while j >= 0 and len(candidates) < 2:
                    if content_lines[j].strip():
                        candidates.append(j)
                    j -= 1
                # The earliest candidate is the entry start
                if candidates:
                    entry_start_indices.add(candidates[-1])
                else:
                    entry_start_indices.add(di)

            # Remove the first entry start (no space needed before first entry)
            if entry_start_indices:
                entry_start_indices.discard(min(entry_start_indices))

            for idx, line in enumerate(content_lines):
                if line.strip():
                    para = doc.add_paragraph()

                    if idx in entry_start_indices:
                        para.paragraph_format.space_before = Pt(10)
                    else:
                        para.paragraph_format.space_before = Pt(2)

                    para.paragraph_format.space_after = Pt(2)
                    run = para.add_run(line.strip())
                    run.font.size = Pt(11)

    # Add ATS score footer
    doc.add_paragraph()  # spacer
    footer_sep = doc.add_paragraph()
    sep_run = footer_sep.add_run("─" * 70)
    sep_run.font.size = Pt(4)
    sep_run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)

    # Score headline
    score_para = doc.add_paragraph()
    score_para.paragraph_format.space_after = Pt(2)
    score_run = score_para.add_run(
        f"ATS Readability Score: {score_data['score']}/100"
    )
    score_run.font.size = Pt(10)
    score_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    score_run.bold = True

    # Passed checks
    if score_data.get("passed"):
        passed_para = doc.add_paragraph()
        passed_para.paragraph_format.space_before = Pt(4)
        passed_para.paragraph_format.space_after = Pt(1)
        passed_label = passed_para.add_run("✓ Passed: ")
        passed_label.font.size = Pt(9)
        passed_label.font.color.rgb = RGBColor(0x4A, 0x67, 0x41)
        passed_label.bold = True
        passed_text = passed_para.add_run(
            " · ".join(score_data["passed"])
        )
        passed_text.font.size = Pt(9)
        passed_text.font.color.rgb = RGBColor(0x4A, 0x67, 0x41)

    # Issues found
    if score_data.get("issues"):
        issues_para = doc.add_paragraph()
        issues_para.paragraph_format.space_before = Pt(2)
        issues_para.paragraph_format.space_after = Pt(1)
        issues_label = issues_para.add_run("✗ To improve: ")
        issues_label.font.size = Pt(9)
        issues_label.font.color.rgb = RGBColor(0xC4, 0x55, 0x3A)
        issues_label.bold = True
        issues_text = issues_para.add_run(
            " · ".join(score_data["issues"])
        )
        issues_text.font.size = Pt(9)
        issues_text.font.color.rgb = RGBColor(0xC4, 0x55, 0x3A)
    else:
        perfect_para = doc.add_paragraph()
        perfect_para.paragraph_format.space_before = Pt(2)
        perfect_para.paragraph_format.space_after = Pt(1)
        perfect_run = perfect_para.add_run(
            "No issues found — your resume is fully ATS-optimised."
        )
        perfect_run.font.size = Pt(9)
        perfect_run.font.color.rgb = RGBColor(0x4A, 0x67, 0x41)

    # Branding
    brand_para = doc.add_paragraph()
    brand_para.paragraph_format.space_before = Pt(6)
    brand_run = brand_para.add_run("Generated by ATSReady.co.uk")
    brand_run.font.size = Pt(8)
    brand_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    brand_run.italic = True

    doc.save(output_path)


def convert_resume(input_pdf: str, output_docx: str) -> dict:
    """
    Main conversion pipeline.
    Returns a JSON-serialisable summary dict.
    """
    # 1. Extract text
    raw_text = extract_text_from_pdf(input_pdf)

    if not raw_text.strip():
        return {
            "success": False,
            "error": "Could not extract any text from this PDF. "
                     "It may be an image-only PDF. Please ensure your "
                     "resume contains selectable text.",
            "score": 0,
        }

    # 2. Detect sections
    sections = detect_sections(raw_text)

    # 3. Score ATS readability
    score_data = score_ats_readability(raw_text, sections)

    # 4. Build clean DOCX
    build_ats_docx(sections, output_docx, score_data)

    return {
        "success": True,
        "score": score_data["score"],
        "issues": score_data["issues"],
        "passed": score_data["passed"],
        "word_count": score_data["word_count"],
        "sections_detected": len([s for s in sections if s["heading"] != "HEADER"]),
        "output_file": output_docx,
    }


# ─── CLI Entry Point ──────────────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python ats_converter.py <input.pdf> <output.docx>")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    if not os.path.exists(input_path):
        print(json.dumps({"success": False, "error": f"File not found: {input_path}"}))
        sys.exit(1)

    result = convert_resume(input_path, output_path)
    print(json.dumps(result, indent=2))

    sys.exit(0 if result["success"] else 1)
